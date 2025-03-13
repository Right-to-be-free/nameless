import os
from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
from docx import Document
from fpdf import FPDF
from PyPDF2 import PdfMerger, PdfReader
from PIL import Image

app = Flask(__name__)

# Define paths
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "pdfs")

# Ensure folders exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER
ALLOWED_EXTENSIONS = {"txt", "docx", "jpeg", "jpg", "png", "pdf"}

def allowed_file(filename):
    """Check if file extension is allowed."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_text(text):
    """Replace unsupported characters with safe alternatives."""
    replacements = {
        "\u2014": "-", "\u2013": "-", "\u2026": "...",
        "\u201C": '"', "\u201D": '"', "\u2018": "'", "\u2019": "'"
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text

class IEEE_PDF(FPDF):
    """Custom PDF class for IEEE-style formatting."""
    def header(self):
        self.set_font("Arial", "B", 14)
        self.cell(0, 10, "", ln=True, align="C")
        self.ln(10)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")

    def add_section(self, title):
        self.set_font("Arial", "B", 12)
        self.cell(0, 10, title, ln=True, align="L")
        self.ln(5)

    def add_paragraph(self, text):
        self.set_font("Arial", "", 10)
        self.multi_cell(0, 6, text, align="J")
        self.ln(3)

def convert_to_pdf(input_path, output_path):
    """Convert different file types to a formatted PDF."""
    try:
        pdf = IEEE_PDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        pdf.add_section("")
        pdf.add_paragraph("")

        pdf.add_section("")

        if input_path.endswith(".txt"):
            with open(input_path, "r", encoding="utf-8") as file:
                content = file.read()
            pdf.add_paragraph(clean_text(content))

        elif input_path.endswith(".docx"):
            doc = Document(input_path)
            content = "\n".join([para.text for para in doc.paragraphs])
            pdf.add_paragraph(clean_text(content))

        elif input_path.endswith((".jpeg", ".jpg", ".png")):
            pdf.add_section("")
            pdf.image(input_path, x=10, w=180)

        elif input_path.endswith(".pdf"):
            merger = PdfMerger()
            merger.append(input_path)
            merger.write(output_path)
            merger.close()
            return

        pdf.add_section("")
        pdf.add_paragraph("")

        pdf.output(output_path, "F")

        if not os.path.exists(output_path):
            raise Exception("PDF file was not created successfully.")

    except Exception as e:
        print(f"Error during PDF conversion: {e}")

def merge_pdfs(pdf_list, output_path):
    """Merge multiple PDFs into one."""
    merger = PdfMerger()
    for pdf in pdf_list:
        merger.append(pdf)
    merger.write(output_path)
    merger.close()

@app.route("/", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":
        if "files" not in request.files:
            return "No file part"

        files = request.files.getlist("files")
        if not files or files[0].filename == "":
            return "No selected files"

        pdf_files = []
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                input_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                file.save(input_path)

                output_filename = f"{os.path.splitext(filename)[0]}.pdf"
                output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)

                convert_to_pdf(input_path, output_path)
                pdf_files.append(output_path)

        # Merge PDFs if multiple are uploaded
        if len(pdf_files) > 1:
            merged_pdf_path = os.path.join(app.config["OUTPUT_FOLDER"], "merged_output.pdf")
            merge_pdfs(pdf_files, merged_pdf_path)
            return send_file(merged_pdf_path, as_attachment=True)

        return send_file(pdf_files[0], as_attachment=True)

    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
